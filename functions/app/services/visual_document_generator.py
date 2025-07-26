# app/services/visual_document_generator.py
import json
import re
from typing import Dict, List, Optional
from dataclasses import dataclass
from pathlib import Path
import base64
import hashlib
from dotenv import load_dotenv
import os

load_dotenv()

try:
    from google.cloud import aiplatform
    from google.cloud.aiplatform.gapic.schema import predict
    VERTEX_AI_AVAILABLE = True
except ImportError:
    VERTEX_AI_AVAILABLE = False
    print("Warning: google-cloud-aiplatform not installed. Image generation will be disabled.")

from docx import Document
from docx.shared import Inches
from langchain_google_genai import ChatGoogleGenerativeAI

@dataclass
class ImageRequirement:
    section: str
    description: str
    prompt: str
    image_path: Optional[str] = None

class VisualDocumentGenerator:
    def __init__(self, gemini_api_key: str, project_id: str, location: str = "us-central1"):
        self.gemini_api_key = gemini_api_key
        self.project_id = project_id
        self.location = location
        
        if not VERTEX_AI_AVAILABLE:
            raise ImportError("google-cloud-aiplatform package is required for image generation.")
        
        # Initialize Vertex AI
        aiplatform.init(project=project_id, location=location)
        
        # Initialize LLM for text processing
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-1.5-pro",
            google_api_key=gemini_api_key
        )
        
        # Initialize Vertex AI client for image generation
        self.prediction_client = aiplatform.gapic.PredictionServiceClient()
        
        # Vertex AI Image Generation endpoint
        # Replace with your actual deployed model endpoint
        # self.image_endpoint = f"projects/{project_id}/locations/{location}/endpoints/YOUR_ENDPOINT_ID"
        self.image_endpoint = f"projects/{self.project_id}/locations/{self.location}/publishers/google/models/imagen-4.0-generate-preview-06-06"
    
    def extract_image_requirements(self, lesson_plan: str) -> List[ImageRequirement]:
        """Extract sections that need visual content using multiple approaches"""
        requirements = []
        
        # Method 1: Rule-based extraction
        rule_based_requirements = self._extract_by_rules(lesson_plan)
        requirements.extend(rule_based_requirements)
        
        # Method 2: LLM-based extraction  
        llm_requirements = self._extract_by_llm(lesson_plan)
        requirements.extend(llm_requirements)
        
        # Remove duplicates and return
        unique_requirements = self._remove_duplicates(requirements)
        return unique_requirements
    
    def _extract_by_rules(self, lesson_plan: str) -> List[ImageRequirement]:
        """Rule-based extraction for common visual content patterns"""
        requirements = []
        
        # Common patterns that indicate visual content needs
        visual_patterns = [
            (r"Show a picture of (.+?)(?:\.|Ask|Show)", "Hook Activity"),
            (r"picture story with (.+?)(?:\.|Ask)", "Picture Story"),
            (r"Display (.+?)(?:\.|\")", "Display Material"),
            (r"\(e\.g\.,\s*(.+?)\)", "Example Visual"),
            (r"[Uu]se (.+?pictures?.+?)(?:\.|Use|Show)", "Activity Material"),
            (r"flashcards with (.+?)(?:\.|Play)", "Flashcard Material"),
        ]
        
        for pattern, section_type in visual_patterns:
            matches = re.findall(pattern, lesson_plan, re.IGNORECASE | re.DOTALL)
            for i, match in enumerate(matches):
                description = match.strip()
                if len(description) > 5:
                    requirements.append(ImageRequirement(
                        section=f"{section_type} {i+1}",
                        description=description,
                        prompt=self._generate_image_prompt(description)
                    ))
        
        return requirements
    
    def _extract_by_llm(self, lesson_plan: str) -> List[ImageRequirement]:
        """LLM-based extraction with improved prompting"""
        extraction_prompt = f"""
        Analyze this lesson plan and identify sections that need visual content.

        LESSON PLAN:
        {lesson_plan}

        Find mentions of:
        - Pictures to show students
        - Visual aids or materials
        - Examples that can be illustrated
        - Activities requiring images

        Return ONLY a JSON array like this:
        [
            {{"section": "Hook Activity", "description": "festival scene"}},
            {{"section": "Teaching Example", "description": "child with ice cream"}}
        ]
        """
        
        try:
            response = self.llm.invoke([{"role": "user", "content": extraction_prompt}])
            content = response.content.strip()
            
            print(f"Raw LLM response: {content[:200]}...")
            
            json_pattern = r'\[[\s\S]*?\]'
            match = re.search(json_pattern, content)
            
            if match:
                json_str = match.group(0)
                try:
                    requirements_data = json.loads(json_str)
                    
                    if isinstance(requirements_data, list):
                        processed_requirements = []
                        for req in requirements_data:
                            if isinstance(req, dict) and "section" in req and "description" in req:
                                section = str(req["section"])
                                description = str(req["description"])
                                
                                processed_requirements.append(ImageRequirement(
                                    section=section,
                                    description=description,
                                    prompt=self._generate_image_prompt(description)
                                ))
                        
                        return processed_requirements
                        
                except json.JSONDecodeError as e:
                    print(f"JSON decode error: {e}")
                    
        except Exception as e:
            print(f"LLM extraction failed: {str(e)}")
            
        return []
    
    def _remove_duplicates(self, requirements: List[ImageRequirement]) -> List[ImageRequirement]:
        """Remove duplicate requirements"""
        unique_requirements = []
        seen_descriptions = set()
        
        for req in requirements:
            normalized = req.description.lower().strip()
            if normalized not in seen_descriptions:
                unique_requirements.append(req)
                seen_descriptions.add(normalized)
        
        return unique_requirements
    
    def _generate_image_prompt(self, description: str) -> str:
        """Generate detailed prompts for image generation"""
        if not isinstance(description, str):
            description = str(description)
        
        base_prompt = f"Create an educational illustration showing {description}"
        
        description_lower = description.lower()
        
        if any(word in description_lower for word in ['festival', 'celebration']):
            style_guide = "vibrant festival scene, colorful, joyful, child-friendly cartoon style"
        elif any(word in description_lower for word in ['child', 'student']):
            style_guide = "simple cartoon illustration, bright colors, happy children"
        elif any(word in description_lower for word in ['classroom', 'school']):
            style_guide = "clean classroom setting, educational materials visible"
        else:
            style_guide = "simple, clear educational illustration, bright colors, child-friendly"
        
        enhanced_prompt = f"{base_prompt}. Style: {style_guide}. Educational content for young learners."

        # print(enhanced_prompt)
        
        return enhanced_prompt
    
    def generate_image(self, prompt: str, section_name: str) -> Optional[str]:
        """Generate image using Vertex AI Image Generation"""
        try:
            # print(f"Generating image with prompt: {prompt[:100]}...")
            
            # Prepare the request for Vertex AI
            instances = [
                {
                    "prompt": prompt
                }
            ]
            
            # Optional parameters - adjust based on your model
            parameters = {
                "sampleCount": 1,
                "aspectRatio": "1:1",  # Square images work well for educational content
                "safetyFilterLevel": "block_some",
                "personGeneration": "allow_adult"
            }
            
            # Make the prediction request
            response = self.prediction_client.predict(
                endpoint=self.image_endpoint,
                instances=instances,
                parameters=parameters
            )
            
            # Create output directory
            output_dir = Path("generated_images")
            output_dir.mkdir(exist_ok=True)
            
            # Process the response
            for i, prediction in enumerate(response.predictions):
                # The response structure may vary depending on the model
                # Common formats include:
                # - prediction.get("bytesBase64Encoded") 
                # - prediction.get("image", {}).get("bytesBase64Encoded")
                
                image_data = None
                if "bytesBase64Encoded" in prediction:
                    image_data = prediction["bytesBase64Encoded"]
                elif "image" in prediction and "bytesBase64Encoded" in prediction["image"]:
                    image_data = prediction["image"]["bytesBase64Encoded"]
                
                if image_data:
                    # Decode base64 image
                    image_bytes = base64.b64decode(image_data)
                    
                    # Generate unique filename
                    filename_hash = hashlib.md5(prompt.encode()).hexdigest()[:8]
                    image_path = output_dir / f"{section_name.replace(' ', '_').lower()}_{filename_hash}.png"
                    
                    # Save the image
                    with open(image_path, "wb") as f:
                        f.write(image_bytes)
                    
                    print(f"Generated image saved to: {image_path}")
                    return str(image_path)
            
            print("No image data found in response")
            return None
            
        except Exception as e:
            print(f"Error generating image for {section_name}: {str(e)}")
            # Create a fallback text file for debugging
            output_dir = Path("generated_images")
            output_dir.mkdir(exist_ok=True)
            
            filename_hash = hashlib.md5(prompt.encode()).hexdigest()[:8]
            fallback_path = output_dir / f"{section_name.replace(' ', '_').lower()}_{filename_hash}_error.txt"
            
            with open(fallback_path, 'w') as f:
                f.write(f"Image generation failed\nPrompt: {prompt}\nSection: {section_name}\nError: {str(e)}")
            
            return str(fallback_path)
    
    def create_visual_document(self, lesson_plan: str, images: Dict[str, str]) -> str:
        """Create Word document with integrated images"""
        doc = Document()
        doc.add_heading('Visual Lesson Plan', 0)
        
        sections = self._parse_lesson_sections(lesson_plan)
        
        for section_title, content in sections.items():
            doc.add_heading(section_title, level=1)
            doc.add_paragraph(content)
            
            # Add actual image if available and it's a valid image file
            section_key = section_title.replace(' ', '_').lower()
            print("Sections:", list(sections.keys()))
            print("Image keys:", list(images.keys()))
            if section_key in images and images[section_key]:
                image_path = images[section_key]
                if Path(image_path).exists() and image_path.endswith(('.png', '.jpg', '.jpeg')):
                    try:
                        # Add the actual image to the document
                        doc.add_picture(image_path, width=Inches(4))
                        doc.add_paragraph(f"Generated image for: {section_title}")
                    except Exception as e:
                        doc.add_paragraph(f"[Image generation error: {str(e)}]")
                        doc.add_paragraph(f"Image file: {image_path}")
                else:
                    doc.add_paragraph(f"[Visual Content: {section_title}]")
                    doc.add_paragraph(f"Image file: {image_path}")
                
                doc.add_paragraph()
        
        output_path = "visual_lesson_plan.docx"
        doc.save(output_path)
        return output_path
    
    def _parse_lesson_sections(self, lesson_plan: str) -> Dict[str, str]:
        """Parse lesson plan into sections"""
        sections = {}
        current_section = "Overview"
        current_content = []
        
        for line in lesson_plan.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            # Check for section headers
            if (line.startswith('**') and line.endswith('**') or
                line.startswith('#') or 
                line.endswith(':') or
                any(keyword in line.lower() for keyword in ['hook', 'objective', 'instruction', 'practice'])):
                
                # Save previous section
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                
                current_section = line.strip('*#:').strip()
                current_content = []
            else:
                current_content.append(line)
        
        # Save last section
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections